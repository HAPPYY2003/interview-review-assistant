from __future__ import annotations

import io
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from backend.app.services.text_utils import repair_mojibake


class DocumentParseError(ValueError):
    pass


@dataclass
class ParsedDocument:
    text: str
    metadata: dict[str, object]


class DocumentParser:
    allowed_suffixes = {".txt", ".pdf", ".docx"}

    def __init__(self, max_file_bytes: int):
        self.max_file_bytes = max_file_bytes

    def parse(self, filename: str, content: bytes) -> ParsedDocument:
        if len(content) > self.max_file_bytes:
            raise DocumentParseError("文件超过 5MB 限制")
        suffix = Path(filename or "").suffix.lower()
        if suffix not in self.allowed_suffixes:
            raise DocumentParseError("仅支持 TXT、PDF 和 DOCX 文件")
        if suffix == ".txt":
            text = self._decode_text(content)
            metadata = {"format": "txt", "characters": len(text)}
        elif suffix == ".pdf":
            text, pages = self._parse_pdf(content)
            metadata = {"format": "pdf", "pages": pages, "characters": len(text)}
        else:
            text, paragraphs = self._parse_docx(content)
            metadata = {"format": "docx", "paragraphs": paragraphs, "characters": len(text)}
        text = repair_mojibake(text).strip()
        if not text:
            if suffix == ".pdf":
                raise DocumentParseError("PDF 未提取到文本，可能是扫描件；第一版暂不支持 OCR")
            raise DocumentParseError("文件没有可读取的文本")
        return ParsedDocument(text=text, metadata=metadata)

    @staticmethod
    def _decode_text(content: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise DocumentParseError("TXT 编码无法识别，请转换为 UTF-8")

    @staticmethod
    def _parse_pdf(content: bytes) -> tuple[str, int]:
        try:
            reader = PdfReader(io.BytesIO(content))
            pages = [f"[第 {index} 页]\n{page.extract_text() or ''}" for index, page in enumerate(reader.pages, 1)]
            return "\n\n".join(pages), len(reader.pages)
        except Exception as exc:
            raise DocumentParseError(f"PDF 解析失败：{exc}") from exc

    @staticmethod
    def _parse_docx(content: bytes) -> tuple[str, int]:
        try:
            document = Document(io.BytesIO(content))
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            return "\n".join(f"[段落 {index}] {text}" for index, text in enumerate(paragraphs, 1)), len(paragraphs)
        except Exception as exc:
            raise DocumentParseError(f"DOCX 解析失败：{exc}") from exc
